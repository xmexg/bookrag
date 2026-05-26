from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from .models import LoadedDocument


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def is_supported_file(path: Path, extensions: Iterable[str] | None = None) -> bool:
    allowed = {ext.lower() for ext in (extensions or SUPPORTED_EXTENSIONS)}
    return path.is_file() and path.suffix.lower() in allowed


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return LoadedDocument(text=path.read_text(encoding="utf-8", errors="ignore"), source_type=suffix)

    if suffix == ".pdf":
        return LoadedDocument(text=_load_pdf(path), source_type=suffix)

    if suffix == ".docx":
        return LoadedDocument(text=_load_docx(path), source_type=suffix)

    return LoadedDocument(text=path.read_text(encoding="utf-8", errors="ignore"), source_type="text")


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("读取 PDF 需要安装 pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"[PAGE {page_index}]\n{text.strip()}")
    return "\n\n".join(pages).strip()


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("读取 DOCX 需要安装 python-docx") from exc

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs).strip()
